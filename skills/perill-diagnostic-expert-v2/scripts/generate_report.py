#!/usr/bin/env python3
"""
PERILL Diagnostic Report Generator
Generates comprehensive Word document with embedded charts
Performs qualitative analysis (Layers 5-7) in single AI call
"""

import json
import sys
import base64
import re
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import os

DIMENSION_ORDER = ['P', 'E', 'R', 'I', 'L', 'Ld']
DIMENSION_NAMES = {
    'P': '目的与动机',
    'E': '外部系统与流程',
    'R': '关系',
    'I': '内部系统与流程',
    'L': '学习',
    'Ld': '领导力'
}

DIMENSION_CONTENT_GUIDANCE = {
    'P': 'Focus on mission clarity, goal alignment, value connection',
    'E': 'Focus on stakeholder relations, resource access, market awareness',
    'R': 'Focus on trust, psychological safety, conflict handling',
    'I': 'Focus on processes, decision-making, role clarity, communication',
    'L': 'Focus on reflection habits, development plans, feedback culture',
    'Ld': 'Focus on leadership distribution, leader-team dynamics'
}

# Health assessment
HEALTH_THRESHOLDS = [
    (4.0, 5.0, '高绩效'),
    (3.5, 4.0, '相对健康'),
    (3.0, 3.5, '中等风险'),
    (2.0, 3.0, '核心瓶颈'),
    (1.0, 2.0, '危机')
]

GAP_THRESHOLDS = [
    (1.5, float('inf'), '极强'),
    (1.0, 1.5, '强'),
    (0.5, 1.0, '中等'),
    (0.0, 0.5, '弱')
]

INTERCONNECTION_PAIRS = [
    ('E', 'I', '外部协作与内部流程：输入-执行闭环'),
    ('R', 'I', '关系信任与内部授权：安全-责任闭环'),
    ('I', 'L', '内部流程与学习反思：执行-复盘闭环'),
    ('P', 'R', '共同目的与关系信任：愿景-凝聚闭环'),
    ('Ld', 'I', '领导风格与内部授权：控制-依赖闭环'),
    ('E', 'L', '外部市场与学习发展：需求-适应闭环')
]


def get_health_level(score: float) -> str:
    for lo, hi, level in HEALTH_THRESHOLDS:
        if lo <= score <= hi:
            return level
    return '未知'


def get_gap_level(gap: float) -> str:
    for lo, hi, level in GAP_THRESHOLDS:
        if lo <= gap < hi or (gap >= lo and hi == float('inf')):
            return level
    return '弱'


def load_analysis_json(json_path: str) -> Dict:
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


def load_chart_as_base64(html_path: str) -> Optional[str]:
    """Convert HTML chart to PNG using chromium"""
    html_path = Path(html_path)
    if not html_path.exists():
        return None

    # Try using kaleido from plotly if available
    try:
        import plotly.io as pio
        # Read HTML and extract figure data
        with open(html_path, 'r', encoding='utf-8') as f:
            html_content = f.read()

        # Try to extract JSON config from HTML
        import re
        match = re.search(r'Plot\.newPlot\(.*?\{([\s\S]*?)\}\)', html_content)
        if not match:
            return None

        # Use kaleodo to save - but we need the figure object
        # Since we don't have the figure, we'll use selenium or webkit
        return None  # Fallback to no image
    except:
        return None


def html_to_image(html_path: str, output_path: str) -> bool:
    """Convert HTML to PNG using selenium or similar"""
    html_path = Path(html_path)
    if not html_path.exists():
        return False

    # Try using chromium via subprocess
    try:
        # Check if chromium is available
        result = subprocess.run(
            ['which', 'chromium-browser'] + ['||'] + ['which', 'google-chrome'] + ['||'] + ['which', 'chrome'],
            shell=True,
            capture_output=True
        )
        # For now, just return False - image will be referenced by path
        return False
    except:
        return False


def add_heading(doc: Document, text: str, level: int = 1):
    """Add styled heading"""
    heading = doc.add_heading(text, level=level)
    return heading


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False):
    """Add styled paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def add_bullet_list(doc: Document, items: List[str], indent: int = 0):
    """Add bullet list"""
    for item in items:
        para = doc.add_paragraph(item, style='List Bullet')
        if indent > 0:
            para.paragraph_format.left_indent = Inches(0.25 * indent)


def create_dimension_table(doc: Document, data: Dict) -> None:
    """Create Section 1.1 diagnostic snapshot table"""
    table = doc.add_table(rows=7, cols=6)
    table.style = 'Table Grid'

    # Header row
    headers = ['维度', '团队现状', '相关方现状', '健康评估', '差距分析', '提升需求']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True

    # Data rows
    l1 = data['layer1']
    l3 = data['layer3']
    l4 = data['layer4']

    for i, dim in enumerate(DIMENSION_ORDER):
        ds = l1['dimension_scores'][dim]
        row = table.rows[i + 1]

        row.cells[0].text = f"{dim} ({DIMENSION_NAMES[dim]})"
        row.cells[1].text = f"{ds['current']['team']:.2f}"
        row.cells[2].text = f"{ds['current']['stakeholder']:.2f}"
        row.cells[3].text = l3['health'][dim]['assessment']
        row.cells[4].text = f"{l4['gaps'][dim]['gap']:.2f}"
        row.cells[5].text = l4['improvements'][dim]['level']


def embed_image(doc: Document, html_path: Path, width: Inches = Inches(5.5)):
    """Embed chart from HTML file as image"""
    html_path = Path(html_path)
    if not html_path.exists():
        # Add placeholder text
        doc.add_paragraph(f"[Chart not available: {html_path.name}]")
        return

    # Try to convert HTML to image using a simple approach
    # For now, we'll use a reference path approach
    # In production, you'd use selenium/webkit or kaleido

    # Add placeholder indicating chart location
    doc.add_paragraph(f"[Figure: {html_path.name}]")


def build_ai_analysis_prompt(data: Dict, team_background: str) -> str:
    """Build comprehensive prompt for AI qualitative analysis"""

    l1 = data['layer1']
    l2 = data['layer2']
    l3 = data['layer3']
    l4 = data['layer4']
    raw = data['raw_data']

    # Extract key metrics for prompt
    dimension_metrics = []
    for dim in DIMENSION_ORDER:
        ds = l1['dimension_scores'][dim]
        dim_metrics = {
            'abbr': dim,
            'name_zh': DIMENSION_NAMES[dim],
            'team_current': round(ds['current']['team'], 2),
            'sth_current': round(ds['current']['stakeholder'], 2),
            'team_expected': round(ds['expected']['team'], 2),
            'gap': round(l4['gaps'][dim]['gap'], 2),
            'health': l3['health'][dim]['assessment'],
            'blindspot_gap': round(l3['blindspots'][dim]['gap'], 2),
            'improvement_level': l4['improvements'][dim]['level']
        }
        dimension_metrics.append(dim_metrics)

    # Top/bottom questions
    top_qs = l2['team_top4']
    bottom_qs = l2['team_bottom4']

    # Question details for each dimension
    q_table = []
    for q in raw['question_scores_table']:
        q_table.append({
            'q_num': q['q_num'],
            'dim': q['dimension'],
            'q_zh': q['q_zh'],
            'current': round(q['current_team'], 2),
            'expected': round(q['expected_team'], 2)
        })

    # Build prompt
    prompt = f"""## PERILL Team Diagnostic — Qualitative Analysis (Layers 5-7)

You are a team coaching expert using the PERILL diagnostic model. Generate comprehensive qualitative analysis for a team diagnostic report.

### Team Background
{team_background if team_background else '[No team background provided]'}

### Quantitative Data Summary

#### Dimension Scores
| 维度 | 团队现状 | 相关方现状 | 健康评估 | 差距 | 提升需求 |
|------|---------|-----------|---------|------|---------|
"""

    for dm in dimension_metrics:
        prompt += f"| {dm['abbr']} ({dm['name_zh']}) | {dm['team_current']} | {dm['sth_current']} | {dm['health']} | {dm['gap']} | {dm['improvement_level']} |\n"

    prompt += f"""
#### Top/Bottom Questions (Team Self-Assessment)
**Strengths:**
"""
    for q in top_qs:
        prompt += f"- Q{q['q_num']} ({q['dimension']}): {q['q_text']} = {q['score']}\n"

    prompt += "**Weaknesses:**\n"
    for q in bottom_qs:
        prompt += f"- Q{q['q_num']} ({q['dimension']}): {q['q_text']} = {q['score']}\n"

    prompt += """
#### Question-Level Details
```
"""

    for q in q_table:
        prompt += f"Q{q['q_num']} ({q['dim']}): {q['q_zh']} | 现状={q['current']} 期待={q['expected']}\n"

    prompt += """```
"""

    prompt += """
### Analysis Tasks

Generate ALL of the following in a single comprehensive response:

---

#### LAYER 5: Six-Dimension Deep Diagnosis

For EACH dimension (P, E, R, I, L, Ld), provide:

**{Dimension} — {Chinese Name}**
- **Scores**: Team Current / Stakeholder Current / Expected

**Advantage Evidence — Team Background** (2 bullet points extracted from background)
**Advantage Evidence — Questionnaire Data** (2 bullet points with specific question scores, e.g., Q15=4.00)
**Advantage Analysis** (2-3 sentences):
  - Sentence 1: Identify the team's specific strengths in this dimension connected to evidence
  - Sentence 2+: Explain the positive significance for the team's future development as a system

**Problem Evidence — Team Background** (3 bullet points from background)
**Problem Evidence — Questionnaire Data** (3-4 bullet points with specific question scores, e.g., Q13=2.50)
**Core Problem Summary** (2-3 sentences, specific to team's situation)
**System Impact Analysis** (2-3 sentences about impact on team's future development as a system)
**Diagnostic Conclusion** (10 words or less, in「」brackets)

---

#### LAYER 6: Interconnection Analysis

**6.1 Dimension Interconnection Analysis**
For EACH of the 6 dimension pairs, provide in flowing prose (NOT bullets):

```
Dimension Pair: {e.g., E ↔ I}
Causal Overview: {brief description}

[Analysis text explaining: closed-loop problem → team system impact → intervention approach]
```

Pairs to analyze:
1. E ↔ I (外部协作与内部流程：输入-执行闭环)
2. R ↔ I (关系信任与内部授权：安全-责任闭环)
3. I ↔ L (内部流程与学习反思：执行-复盘闭环)
4. P ↔ R (共同目的与关系信任：愿景-凝聚闭环)
5. Ld ↔ I (领导风格与内部授权：控制-依赖闭环)
6. E ↔ L (外部市场与学习发展：需求-适应闭环)

**6.2 Current Causal Cycle**
- Write 1 paragraph describing the current negative reinforcement cycle
- Use ONLY PERILL dimension abbreviations (P, E, R, I, L, Ld) — NO question numbers
- Example format: "领导力缺乏（Ld）→ 内部系统僵化（I）→ ..."

**6.3 Intervention Causal Cycle**
- Write 1 paragraph describing the positive intervention cycle
- Use dimension abbreviations with brief action descriptions
- Example format: "从关系破冰（R）→ 建立心理安全感 → 内部授权（I）→ ..."

---

#### LAYER 7: Intervention Design

**7.1 Core Diagnostic Conclusion**
MUST be exactly 3 logically connected sentences:
1. Sentence 1: Describe lever dimension(s) and goal of intervention strategy
2. Sentence 2: Describe expected state across all dimensions after intervention (use PERILL abbreviations)
3. Sentence 3: Insightful concluding statement (thought-provoking golden sentence)

**7.2 Lever Dimension Selection**
Justify chosen lever dimension(s) based on:
- Lowest score
- Highest gap
- Strongest interconnections
- Coachability

**7.3 Phase-by-Phase Intervention Strategy**
Table format (4-5 phases):
| Phase | Name (5-10 chars) | Team System Goal | Key Actions | Expected Team System Output |
|-------|-------------------|-----------------|-------------|----------------------------|
| 1 | ... | ... | ... | ... |

IMPORTANT: Goals must focus on TEAM AS A SYSTEM performance, NOT individual leader performance.
- ✅ Correct: "提升团队整体的运作效率"
- ❌ Wrong: "团队Leader的领导力提升"

**7.4 Expected Outcomes and Milestones**
6-month team system level outcomes for each dimension.

---

### Output Format
Return your analysis in structured JSON format:

```json
{{
  "layer5": {{
    "P": {{ "scores": {{...}}, "advantage_bg": [...], "advantage_data": [...], "advantage_analysis": "...", "problem_bg": [...], "problem_data": [...], "core_problem": "...", "system_impact": "...", "conclusion": "..." }},
    "E": {{ ... }},
    "R": {{ ... }},
    "I": {{ ... }},
    "L": {{ ... }},
    "Ld": {{ ... }}
  }},
  "layer6": {{
    "interconnections": [
      {{ "pair": "E ↔ I", "causal_overview": "...", "analysis": "..." }},
      ...
    ],
    "current_cycle": {{
      "name": "...",
      "path": "...",
      "characteristics": "..."
    }},
    "intervention_cycle": {{
      "name": "...",
      "path": "...",
      "characteristics": "..."
    }}
  }},
  "layer7": {{
    "core_conclusion": {{
      "sentence1": "...",
      "sentence2": "...",
      "sentence3": "..."
    }},
    "lever_dimensions": [...],
    "lever_justification": "...",
    "phases": [
      {{ "phase": 1, "name": "...", "team_goal": "...", "key_actions": "...", "expected_output": "..." }},
      ...
    ],
    "expected_outcomes": {{...}}
  }}
}}
```

Generate the complete analysis now.
"""

    return prompt


def call_ai_for_qualitative_analysis(prompt: str, api_key: str, api_base: str) -> str:
    """Call AI API for qualitative analysis"""
    import urllib.request
    import urllib.error

    payload = {
        "model": "gpt-4o",
        "messages": [
            {"role": "system", "content": "You are an expert team coaching consultant using the PERILL diagnostic model. Provide thorough, insightful analysis that is grounded in the data provided."},
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.3,
        "max_tokens": 8000
    }

    data = json.dumps(payload).encode('utf-8')
    headers = {
        'Content-Type': 'application/json',
        'Authorization': f'Bearer {api_key}'
    }

    req = urllib.request.Request(api_base, data=data, headers=headers, method='POST')

    try:
        with urllib.request.urlopen(req, timeout=120) as response:
            result = json.loads(response.read().decode('utf-8'))
            return result['choices'][0]['message']['content']
    except urllib.error.HTTPError as e:
        error_body = e.read().decode('utf-8') if e.fp else ''
        raise Exception(f"AI API error {e.code}: {error_body}")
    except Exception as e:
        raise Exception(f"AI API call failed: {str(e)}")


def parse_ai_response(response: str) -> Dict:
    """Parse AI response, extracting JSON from markdown code blocks"""
    # Try to find JSON in code blocks
    import re

    # Look for JSON in code blocks
    json_match = re.search(r'```json\s*([\s\S]*?)\s*```', response)
    if json_match:
        json_str = json_match.group(1)
    else:
        # Try to find any JSON object
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            json_str = json_match.group(0)
        else:
            raise Exception("Could not find JSON in AI response")

    # Parse JSON
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        raise Exception(f"Failed to parse JSON: {e}\nContent: {json_str[:500]}")


def generate_report(analysis_json_path: str, team_background: str, output_docx_path: str,
                    api_key: str = None, api_base: str = None) -> str:
    """
    Generate complete Word document report
    """

    # Load analysis data
    data = load_analysis_json(analysis_json_path)
    raw = data['raw_data']

    # Get API credentials from config if not provided
    if not api_key or not api_base:
        config_path = Path(__file__).parent.parent / 'config.json'
        if config_path.exists():
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                api_key = api_key or config.get('api_key', '')
                api_base = api_base or config.get('api_base', 'https://api.openai.com/v1/chat/completions')

    # Create document
    doc = Document()

    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ===== DOCUMENT TITLE =====
    title = doc.add_heading('PERILL团队诊断报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph('PERILL Team Diagnostic Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Report metadata
    from datetime import datetime
    date_str = datetime.now().strftime('%Y年%m月%d日')

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f'诊断日期: {date_str}\n').bold = True
    meta.add_run('诊断模型: PERILL\n')
    meta.add_run(f'团队成员数: {data["metadata"]["respondent_count"]}\n')

    respondents = data['metadata']['respondents']
    team_members = [r for r in respondents if r['role'] in ['Team Member', 'Team Leader']]
    stakeholders = [r for r in respondents if r['role'] == 'Stakeholder']

    meta.add_run(f'答卷人数: {len(respondents)} (团队成员: {len(team_members)}, 相关方: {len(stakeholders)})')

    doc.add_paragraph()  # Spacing

    # ===== SECTION 1: DIAGNOSTIC OVERVIEW =====
    add_heading(doc, '一、诊断概览', 1)

    add_heading(doc, '1.1 PERILL六维诊断快照', 2)

    create_dimension_table(doc, data)

    doc.add_paragraph()  # Spacing

    # Embedded visualizations
    add_heading(doc, '现状与期望对比', 3)
    charts_dir = Path(analysis_json_path).parent / 'charts'

    embed_image(doc, charts_dir / 'radar_current_vs_expected.html')
    doc.add_paragraph()

    add_heading(doc, '团队自评与利益相关者评价对比', 3)
    embed_image(doc, charts_dir / 'radar_team_vs_stakeholder.html')
    doc.add_paragraph()

    add_heading(doc, '1.2 核心发现摘要', 2)

    # Core findings from quantitative data
    l4 = data['layer4']
    l3 = data['layer3']

    # Find lowest scoring and highest gap dimensions
    dim_gaps = [(dim, l4['gaps'][dim]['gap']) for dim in DIMENSION_ORDER]
    dim_gaps.sort(key=lambda x: x[1], reverse=True)

    add_bullet_list(doc, [
        f"最低得分维度: {DIMENSION_ORDER[raw['dimension_current_team'].index(min(raw['dimension_current_team']))]} ({DIMENSION_NAMES[DIMENSION_ORDER[raw['dimension_current_team'].index(min(raw['dimension_current_team']))]]}) = {min(raw['dimension_current_team']):.2f})",
        f"最大提升需求: {dim_gaps[0][0]} ({DIMENSION_NAMES[dim_gaps[0][0]]}), 差距 = {dim_gaps[0][1]:.2f}",
        f"存在认知盲区: {[dim for dim in DIMENSION_ORDER if '盲区' in l3['blindspots'][dim]['blindspot_type']]}"
    ])

    doc.add_paragraph()

    add_heading(doc, '1.3 核心诊断结论', 2)

    # Check if we have AI qualitative analysis
    ai_qualitative_path = Path(analysis_json_path).parent / 'qualitative_analysis.json'
    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)
        conclusion = qualitative.get('layer7', {}).get('core_conclusion', {})
        if conclusion:
            p1 = doc.add_paragraph()
            p1.add_run(conclusion.get('sentence1', '[待AI分析生成]'))

            p2 = doc.add_paragraph()
            p2.add_run(conclusion.get('sentence2', ''))

            p3 = doc.add_paragraph()
            p3.add_run(conclusion.get('sentence3', ''))
        else:
            doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    doc.add_page_break()

    # ===== SECTION 2: SIX-DIMENSION DEEP DIAGNOSIS =====
    add_heading(doc, '二、六维深度诊断', 1)

    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)

        layer5 = qualitative.get('layer5', {})

        for dim in DIMENSION_ORDER:
            dim_data = layer5.get(dim, {})
            if not dim_data:
                continue

            add_heading(doc, f'{dim} — {DIMENSION_NAMES[dim]}', 2)

            scores = dim_data.get('scores', {})
            doc.add_paragraph(f"得分: 团队现状 {scores.get('team_current', 'N/A')} / 相关方现状 {scores.get('sth_current', 'N/A')} / 期望 {scores.get('expected', 'N/A')}")

            doc.add_paragraph().add_run('优势分析').bold = True
            add_heading(doc, '优势证据 — 团队背景', 3)
            add_bullet_list(doc, dim_data.get('advantage_bg', ['[待分析]']))

            add_heading(doc, '优势证据 — 问卷数据', 3)
            add_bullet_list(doc, dim_data.get('advantage_data', ['[待分析]']))

            add_heading(doc, '优势分析', 3)
            doc.add_paragraph(dim_data.get('advantage_analysis', '[待分析]'))

            doc.add_paragraph().add_run('问题分析').bold = True
            add_heading(doc, '问题证据 — 团队背景', 3)
            add_bullet_list(doc, dim_data.get('problem_bg', ['[待分析]']))

            add_heading(doc, '问题证据 — 问卷数据', 3)
            add_bullet_list(doc, dim_data.get('problem_data', ['[待分析]']))

            add_heading(doc, '核心问题', 3)
            doc.add_paragraph(dim_data.get('core_problem', '[待分析]'))

            add_heading(doc, '系统影响', 3)
            doc.add_paragraph(dim_data.get('system_impact', '[待分析]'))

            add_heading(doc, '诊断结论', 3)
            conclusion_para = doc.add_paragraph()
            conclusion_para.add_run(f"「{dim_data.get('conclusion', '[待分析]')}」").bold = True

            doc.add_paragraph()  # Spacing
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    doc.add_page_break()

    # ===== SECTION 3: SYSTEMIC INTERCONNECTION =====
    add_heading(doc, '三、系统性关联分析', 1)

    add_heading(doc, '3.1 维度关联分析', 2)

    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)

        interconnections = qualitative.get('layer6', {}).get('interconnections', [])

        for ic in interconnections:
            pair = ic.get('pair', '')
            causal = ic.get('causal_overview', '')
            analysis = ic.get('analysis', '')

            p = doc.add_paragraph()
            p.add_run(f"维度对: {pair}\n").bold = True
            p.add_run(f"因果概述: {causal}\n\n")
            p.add_run(analysis)
            doc.add_paragraph()
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    add_heading(doc, '3.2 现状因果循环', 2)

    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)

        current_cycle = qualitative.get('layer6', {}).get('current_cycle', {})
        if current_cycle:
            p = doc.add_paragraph()
            p.add_run(f"循环名称: {current_cycle.get('name', '待分析')}\n\n").bold = True
            p.add_run(f"循环构成路径:\n{current_cycle.get('path', '')}\n\n")
            p.add_run(current_cycle.get('characteristics', ''))
        else:
            doc.add_paragraph('[待分析]')
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    add_heading(doc, '3.3 干预因果循环', 2)

    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)

        intervention_cycle = qualitative.get('layer6', {}).get('intervention_cycle', {})
        if intervention_cycle:
            p = doc.add_paragraph()
            p.add_run(f"循环名称: {intervention_cycle.get('name', '待分析')}\n\n").bold = True
            p.add_run(f"循环构成路径:\n{intervention_cycle.get('path', '')}\n\n")
            p.add_run(intervention_cycle.get('characteristics', ''))
        else:
            doc.add_paragraph('[待分析]')
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    add_heading(doc, '3.4 系统因果网络图', 2)
    embed_image(doc, charts_dir / 'system_network.html')

    doc.add_page_break()

    # ===== SECTION 4: INTERVENTION PATH =====
    add_heading(doc, '四、干预路径建议', 1)

    add_heading(doc, '4.1 核心诊断结论', 2)

    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)

        conclusion = qualitative.get('layer7', {}).get('core_conclusion', {})
        if conclusion:
            doc.add_paragraph(conclusion.get('sentence1', ''))
            doc.add_paragraph(conclusion.get('sentence2', ''))
            doc.add_paragraph(conclusion.get('sentence3', ''))
        else:
            doc.add_paragraph('[待分析]')
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    add_heading(doc, '4.2 杠杆维度选择', 2)

    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)

        lever_dims = qualitative.get('layer7', {}).get('lever_dimensions', [])
        lever_justification = qualitative.get('layer7', {}).get('lever_justification', '')

        if lever_dims:
            p = doc.add_paragraph()
            p.add_run('选择杠杆维度: ').bold = True
            p.add_run(', '.join(lever_dims))

            doc.add_paragraph()
            doc.add_paragraph(lever_justification)
        else:
            doc.add_paragraph('[待分析]')
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    add_heading(doc, '4.3 阶段性干预策略', 2)

    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)

        phases = qualitative.get('layer7', {}).get('phases', [])

        if phases:
            # Create phase table
            table = doc.add_table(rows=len(phases) + 1, cols=5)
            table.style = 'Table Grid'

            headers = ['阶段', '名称', '团队系统目标', '关键行动', '预期团队系统产出']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            for i, phase in enumerate(phases):
                row = table.rows[i + 1]
                row.cells[0].text = str(phase.get('phase', i + 1))
                row.cells[1].text = phase.get('name', '')
                row.cells[2].text = phase.get('team_goal', '')
                row.cells[3].text = phase.get('key_actions', '')
                row.cells[4].text = phase.get('expected_output', '')
        else:
            doc.add_paragraph('[待分析]')
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    add_heading(doc, '4.4 预期成果与里程碑', 2)

    if ai_qualitative_path.exists():
        with open(ai_qualitative_path, 'r', encoding='utf-8') as f:
            qualitative = json.load(f)

        outcomes = qualitative.get('layer7', {}).get('expected_outcomes', {})

        if outcomes:
            for dim, outcome in outcomes.items():
                p = doc.add_paragraph()
                p.add_run(f"{dim} ({DIMENSION_NAMES.get(dim, dim)}): ").bold = True
                p.add_run(outcome)
        else:
            doc.add_paragraph('[待分析]')
    else:
        doc.add_paragraph('[AI定性分析尚未生成，请先运行定性分析步骤]')

    doc.add_page_break()

    # ===== SECTION 5: APPENDICES =====
    add_heading(doc, '五、附录', 1)

    add_heading(doc, '5.1 原始评分热力图', 2)
    embed_image(doc, charts_dir / 'respondent_heatmap.html')

    add_heading(doc, '5.2 PERILL模型介绍', 2)

    model_intro = """
PERILL是一个基于20年团队绩效研究的六维团队诊断模型，用于评估团队在以下维度的发展状况:

- P (Purpose): 目的与动机 - 团队使命清晰度、目标一致性、价值连接
- E (External): 外部系统与流程 - 利益相关方关系、资源获取、市场感知
- R (Relationship): 关系 - 信任、心理安全、冲突处理
- I (Internal): 内部系统与流程 - 流程、决策、角色清晰、沟通
- L (Learning): 学习 - 反思习惯、发展计划、反馈文化
- Ld (Leadership): 领导力 - 领导力分布、领导与团队动态

该模型通过问卷评估(36题×1-5分)识别团队的优势、瓶颈和改进机会，为团队教练干预提供诊断依据。
"""
    doc.add_paragraph(model_intro)

    # Save document
    output_path = Path(output_docx_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return str(output_path)


def run_qualitative_analysis(analysis_json_path: str, team_background: str,
                              api_key: str = None, api_base: str = None) -> str:
    """
    Run AI qualitative analysis (Layers 5-7) and save results
    """

    # Load analysis data
    data = load_analysis_json(analysis_json_path)

    # Get API credentials
    if not api_key or not api_base:
        config_path = Path(__file__).parent.parent / 'config.json'
        if config_path.exists():
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                api_key = api_key or config.get('api_key', '')
                api_base = api_base or config.get('api_base', 'https://api.openai.com/v1/chat/completions')

    if not api_key:
        raise Exception("API key not provided and not found in config.json")

    # Build prompt
    prompt = build_ai_analysis_prompt(data, team_background)

    print("\n" + "=" * 60)
    print("Running AI Qualitative Analysis (Layers 5-7)")
    print("=" * 60)
    print("\nThis may take a minute...")

    # Call AI
    response = call_ai_for_qualitative_analysis(prompt, api_key, api_base)

    print("\n✅ AI analysis complete!")

    # Parse response
    try:
        qualitative = parse_ai_response(response)
    except Exception as e:
        print(f"Warning: Failed to parse AI JSON response: {e}")
        # Save raw response
        raw_path = Path(analysis_json_path).parent / 'qualitative_analysis_raw.txt'
        with open(raw_path, 'w', encoding='utf-8') as f:
            f.write(response)
        raise

    # Save qualitative analysis
    output_path = Path(analysis_json_path).parent / 'qualitative_analysis.json'
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(qualitative, f, ensure_ascii=False, indent=2)

    print(f"📁 Qualitative analysis saved to: {output_path}")
    return str(output_path)


def main():
    """CLI entry point"""

    if len(sys.argv) < 3:
        print("""
PERILL Report Generator

Usage:
  Step 1 (Qualitative Analysis): python generate_report.py --analyze <analysis_json> [team_background]
  Step 2 (Generate Word Doc):     python generate_report.py --generate <analysis_json> [team_background]

Examples:
  # Run AI qualitative analysis
  python generate_report.py --analyze output.json "团队成立于2020年..."

  # Generate Word document
  python generate_report.py --generate output.json "团队成立于2020年..."
        """)
        sys.exit(1)

    mode = sys.argv[1]

    if mode == '--analyze':
        if len(sys.argv) < 3:
            print("Error: --analyze requires <analysis_json> path")
            sys.exit(1)
        analysis_json = sys.argv[2]
        team_bg = sys.argv[3] if len(sys.argv) > 3 else ""
        run_qualitative_analysis(analysis_json, team_bg)

    elif mode == '--generate':
        if len(sys.argv) < 3:
            print("Error: --generate requires <analysis_json> path")
            sys.exit(1)
        analysis_json = sys.argv[2]
        team_bg = sys.argv[3] if len(sys.argv) > 3 else ""

        # Determine output path
        json_path = Path(analysis_json)
        output_docx = json_path.parent / f"PERILL_Report_{json_path.stem}.docx"

        generate_report(analysis_json, team_bg, str(output_docx))
        print(f"\n📄 Word document generated: {output_docx}")

    else:
        print(f"Unknown mode: {mode}")
        print("Use --analyze or --generate")
        sys.exit(1)


if __name__ == '__main__':
    main()